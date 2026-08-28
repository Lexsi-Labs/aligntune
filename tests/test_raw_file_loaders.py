import pytest
import tempfile
from pathlib import Path
from datasets import Dataset

from aligntune.data.loaders.pdf_loader import PDFLoader
from aligntune.data.loaders.docx_loader import DocxLoader
from aligntune.data.loaders.markdown_loader import MarkdownLoader
from aligntune.data.loaders.directory_loader import DirectoryLoader
from aligntune.data.loaders.resolver import LoaderResolver


# Fixtures for test files


@pytest.fixture
def sample_pdf_file():
    """Create a simple test PDF file."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        pytest.skip("reportlab not installed")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        pdf_path = f.name

    c = canvas.Canvas(pdf_path, pagesize=letter)

    # Page 1
    c.drawString(100, 750, "Sample PDF Document")
    c.drawString(100, 700, "This is the first page of the PDF.")
    c.drawString(100, 650, "It contains some sample text for testing.")

    c.showPage()

    # Page 2
    c.drawString(100, 750, "Page Two")
    c.drawString(100, 700, "This is the second page.")
    c.drawString(100, 650, "More test content here.")

    c.save()
    yield pdf_path

    # Cleanup
    Path(pdf_path).unlink(missing_ok=True)


@pytest.fixture
def sample_docx_file():
    """Create a simple test DOCX file."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_path = f.name

    doc = Document()
    doc.add_paragraph("Sample DOCX Document")
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_paragraph("This is the third paragraph.")
    doc.save(docx_path)

    yield docx_path

    # Cleanup
    Path(docx_path).unlink(missing_ok=True)


@pytest.fixture
def sample_markdown_file():
    """Create a simple test markdown file."""
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", delete=False, encoding="utf-8"
    ) as f:
        markdown_path = f.name
        f.write("""# Sample Markdown Document

## Introduction

This is the introduction section. It contains some sample text.

## Main Content

This is the main content section.

### Subsection

Some detailed information here.

## Conclusion

This concludes the document.
""")

    yield markdown_path

    # Cleanup
    Path(markdown_path).unlink(missing_ok=True)


@pytest.fixture
def temp_dir():
    """Create a temporary directory."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


# PDF Loader Tests


class TestPDFLoader:
    def test_pdf_loader_initialization(self):
        """Test PDFLoader initialization."""
        loader = PDFLoader("/fake/path.pdf", chunk_size=512, chunk_overlap=0)
        assert loader.path == Path("/fake/path.pdf")
        assert loader.chunk_size == 512
        assert loader.chunk_overlap == 0

    def test_pdf_loader_missing_file(self):
        """Test PDFLoader with non-existent file."""
        loader = PDFLoader("/nonexistent/file.pdf")
        with pytest.raises(FileNotFoundError):
            loader.load_file()

    def test_pdf_loader_single_file(self, sample_pdf_file):
        """Test loading a single PDF file."""
        loader = PDFLoader(sample_pdf_file, chunk_size=None)
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        assert "text" in dataset.column_names
        assert "source" in dataset.column_names
        assert "page" in dataset.column_names

    def test_pdf_loader_with_chunking(self, sample_pdf_file):
        """Test PDF loading with chunking."""
        loader = PDFLoader(sample_pdf_file, chunk_size=100, chunk_overlap=0)
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        # With chunking, all chunks should have page = -1
        assert all(page == -1 for page in dataset["page"])

    def test_pdf_loader_without_chunking(self, sample_pdf_file):
        """Test PDF loading without chunking (page-level)."""
        loader = PDFLoader(sample_pdf_file, chunk_size=None)
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        # Without chunking, should have valid page numbers
        assert all(page >= 0 for page in dataset["page"])

    def test_pdf_loader_chunk_text(self):
        """Test text chunking logic."""
        loader = PDFLoader("/fake/path.pdf", chunk_size=10)
        text = "word1 word2 word3 word4 word5 word6 word7 word8 word9 word10 word11 word12"
        chunks = loader._chunk_text(text)

        assert len(chunks) > 1
        assert all(chunk.strip() for chunk in chunks)


# DOCX Loader Tests


class TestDocxLoader:
    def test_docx_loader_initialization(self):
        """Test DocxLoader initialization."""
        loader = DocxLoader("/fake/path.docx", chunk_size=512)
        assert loader.path == Path("/fake/path.docx")
        assert loader.chunk_size == 512

    def test_docx_loader_missing_file(self):
        """Test DocxLoader with non-existent file."""
        loader = DocxLoader("/nonexistent/file.docx")
        with pytest.raises(FileNotFoundError):
            loader.load_file()

    def test_docx_loader_single_file(self, sample_docx_file):
        """Test loading a single DOCX file."""
        loader = DocxLoader(sample_docx_file, chunk_size=None)
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        assert "text" in dataset.column_names
        assert "source" in dataset.column_names
        assert "paragraph_idx" in dataset.column_names

    def test_docx_loader_paragraph_preservation(self, sample_docx_file):
        """Test that DOCX loader preserves paragraph structure."""
        loader = DocxLoader(sample_docx_file, chunk_size=None)
        dataset = loader.load_file()

        # Should have multiple paragraphs
        assert len(dataset) >= 3
        # Without chunking, should have sequential paragraph indices
        paragraphs = dataset["paragraph_idx"]
        assert all(p >= 0 for p in paragraphs)

    def test_docx_loader_with_chunking(self, sample_docx_file):
        """Test DOCX loading with chunking."""
        loader = DocxLoader(sample_docx_file, chunk_size=50)
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        # With chunking, all paragraphs should have idx = -1
        assert all(p == -1 for p in dataset["paragraph_idx"])


# Markdown Loader Tests


class TestMarkdownLoader:
    def test_markdown_loader_initialization(self):
        """Test MarkdownLoader initialization."""
        loader = MarkdownLoader("/fake/path.md", chunk_by="heading")
        assert loader.path == Path("/fake/path.md")
        assert loader.chunk_by == "heading"

    def test_markdown_loader_missing_file(self):
        """Test MarkdownLoader with non-existent file."""
        loader = MarkdownLoader("/nonexistent/file.md")
        with pytest.raises(FileNotFoundError):
            loader.load_file()

    def test_markdown_loader_single_file(self, sample_markdown_file):
        """Test loading a single markdown file."""
        loader = MarkdownLoader(sample_markdown_file, chunk_by="paragraph")
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        assert "text" in dataset.column_names
        assert "source" in dataset.column_names
        assert "heading" in dataset.column_names
        assert "section" in dataset.column_names

    def test_markdown_loader_heading_based_chunking(self, sample_markdown_file):
        """Test markdown loading with heading-based chunking."""
        loader = MarkdownLoader(sample_markdown_file, chunk_by="heading")
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0
        # Should have multiple sections based on headings
        assert len(set(dataset["heading"])) > 1

    def test_markdown_loader_paragraph_based_chunking(self, sample_markdown_file):
        """Test markdown loading with paragraph-based chunking."""
        loader = MarkdownLoader(sample_markdown_file, chunk_by="paragraph")
        dataset = loader.load_file()

        assert isinstance(dataset, Dataset)
        assert len(dataset) > 0

    def test_markdown_parser_by_heading(self):
        """Test heading-based parsing logic."""
        loader = MarkdownLoader("/fake/path.md")
        content = """# Heading 1
Some text here.

## Heading 2
More text.

### Heading 3
Even more text.
"""
        sections = loader._parse_by_heading(content)
        assert len(sections) >= 2
        assert sections[0]["heading"] == "Heading 1"


# Directory Loader Tests


class TestDirectoryLoader:
    def test_directory_loader_initialization(self):
        """Test DirectoryLoader initialization."""
        loader = DirectoryLoader("/fake/dir", pattern="*.md")
        assert loader.path == Path("/fake/dir")
        assert loader.pattern == "*.md"

    def test_directory_loader_missing_dir(self):
        """Test DirectoryLoader with non-existent directory."""
        loader = DirectoryLoader("/nonexistent/directory")
        with pytest.raises(ValueError):
            loader.load_directory()

    def test_directory_loader_empty_dir(self, temp_dir):
        """Test DirectoryLoader with empty directory."""
        loader = DirectoryLoader(str(temp_dir))
        with pytest.raises(ValueError):
            loader.load_directory()

    def test_directory_loader_supported_formats(self):
        """Test supported formats list."""
        formats = DirectoryLoader.supported_formats()
        assert ".json" in formats
        assert ".csv" in formats
        assert ".pdf" in formats
        assert ".docx" in formats
        assert ".md" in formats

    def test_directory_loader_single_markdown(self, temp_dir, sample_markdown_file):
        """Test loading directory with single markdown file."""
        # Copy markdown file to temp directory
        import shutil
        md_path = temp_dir / "test.md"
        shutil.copy(sample_markdown_file, md_path)

        loader = DirectoryLoader(str(temp_dir), pattern="*.md")
        result = loader.load_directory()

        # Should return a Dataset when all files are raw format
        assert isinstance(result, Dataset) or hasattr(result, "__iter__")
        if isinstance(result, Dataset):
            assert len(result) > 0

    def test_directory_loader_multi_format(self, temp_dir, sample_markdown_file, sample_docx_file):
        """Test loading directory with multiple document types."""
        import shutil

        # Copy files to temp directory
        md_path = temp_dir / "test.md"
        docx_path = temp_dir / "test.docx"
        shutil.copy(sample_markdown_file, md_path)
        shutil.copy(sample_docx_file, docx_path)

        loader = DirectoryLoader(str(temp_dir))
        result = loader.load_directory()

        # Should have data from both files
        if isinstance(result, Dataset):
            assert len(result) > 0
        else:
            assert len(result) > 0

    def test_directory_loader_get_loader_for_file(self):
        """Test _get_loader_for_file routing."""
        loader = DirectoryLoader("/fake/dir")

        # Test routing to different loaders
        pdf_loader = loader._get_loader_for_file(Path("test.pdf"))
        assert isinstance(pdf_loader, PDFLoader)

        docx_loader = loader._get_loader_for_file(Path("test.docx"))
        assert isinstance(docx_loader, DocxLoader)

        md_loader = loader._get_loader_for_file(Path("test.md"))
        assert isinstance(md_loader, MarkdownLoader)

        # Test unsupported format
        unknown_loader = loader._get_loader_for_file(Path("test.xyz"))
        assert unknown_loader is None


# Resolver Tests


class TestLoaderResolver:
    def test_resolver_supported_formats(self):
        """Test resolver supported formats."""
        formats = LoaderResolver.supported_formats()
        assert "pdf" in formats
        assert "docx" in formats
        assert "markdown" in formats
        assert "directory" in formats

    def test_resolver_get_loader_by_format(self):
        """Test getting loader by format name."""
        pdf_loader = LoaderResolver.get_loader_by_format("pdf")
        assert pdf_loader == PDFLoader

        docx_loader = LoaderResolver.get_loader_by_format("docx")
        assert docx_loader == DocxLoader

        md_loader = LoaderResolver.get_loader_by_format("markdown")
        assert md_loader == MarkdownLoader

    def test_resolver_invalid_format(self):
        """Test resolver with invalid format."""
        with pytest.raises(ValueError):
            LoaderResolver.get_loader_by_format("invalid_format")

    def test_resolver_pdf_file(self, sample_pdf_file):
        """Test resolver with PDF file."""
        loader = LoaderResolver.resolve(sample_pdf_file)
        assert isinstance(loader, PDFLoader)

    def test_resolver_docx_file(self, sample_docx_file):
        """Test resolver with DOCX file."""
        loader = LoaderResolver.resolve(sample_docx_file)
        assert isinstance(loader, DocxLoader)

    def test_resolver_markdown_file(self, sample_markdown_file):
        """Test resolver with markdown file."""
        loader = LoaderResolver.resolve(sample_markdown_file)
        assert isinstance(loader, MarkdownLoader)

    def test_resolver_directory(self, temp_dir):
        """Test resolver with directory."""
        loader = LoaderResolver.resolve(str(temp_dir))
        assert isinstance(loader, DirectoryLoader)

    def test_resolver_with_kwargs(self, sample_pdf_file):
        """Test resolver passing kwargs to loader."""
        loader = LoaderResolver.resolve(
            sample_pdf_file,
            chunk_size=256,
            chunk_overlap=10,
        )
        assert isinstance(loader, PDFLoader)
        assert loader.chunk_size == 256
        assert loader.chunk_overlap == 10


# Integration Tests


class TestIntegration:
    def test_load_mixed_documents(self, temp_dir, sample_pdf_file, sample_docx_file, sample_markdown_file):
        """Integration test: load directory with mixed document types."""
        import shutil

        # Copy files to temp directory
        shutil.copy(sample_pdf_file, temp_dir / "test1.pdf")
        shutil.copy(sample_docx_file, temp_dir / "test2.docx")
        shutil.copy(sample_markdown_file, temp_dir / "test3.md")

        loader = DirectoryLoader(str(temp_dir))
        result = loader.load_directory()

        # Should have successfully loaded all files
        if isinstance(result, Dataset):
            assert len(result) > 0
            assert "text" in result.column_names
            assert "source" in result.column_names
        else:
            # DatasetDict case
            assert len(result) > 0

    def test_dataset_format_consistency(self, sample_markdown_file):
        """Test that all loaders return consistent Dataset format."""
        loader = MarkdownLoader(sample_markdown_file)
        dataset = loader.load_file()

        # All datasets should have 'text' and 'source' columns
        assert "text" in dataset.column_names
        assert "source" in dataset.column_names

        # Text should be non-empty strings
        texts = dataset["text"]
        assert all(isinstance(t, str) and len(t) > 0 for t in texts)

    def test_chunking_parameters(self):
        """Test chunking with different parameters."""
        loader = MarkdownLoader("/fake/path.md", chunk_size=100, chunk_overlap=20)

        text = " ".join([f"word{i}" for i in range(100)])
        chunks = loader._chunk_text(text)

        # Should produce multiple chunks
        assert len(chunks) > 1

        # All chunks should be non-empty
        assert all(chunk.strip() for chunk in chunks)
